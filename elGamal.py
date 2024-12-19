import random
from docx import Document
import streamlit as st

# Fungsi untuk membuat kunci El Gamal
def generate_key(p, g, x):
    y = pow(g, x, p)  # y = g^x mod p
    return p, g, y, x

# Fungsi untuk enkripsi El Gamal
def encrypt(p, g, y, k, plaintext):
    C1 = pow(g, k, p)  # C1 = g^k mod p
    ciphertext = []
    space_positions = []

    for i, char in enumerate(plaintext):
        if char == " ":
            space_positions.append(i)
        elif char.isalpha():
            base = ord('A') if char.isupper() else ord('a')
            M = ord(char) - base  # Konversi karakter ke angka 0-25 (atau 0-25 untuk huruf kecil)
            y_k = pow(y, k, p)  # Hitung y^k mod p
            C2 = (M * y_k) % p  # C2 = M * y^k mod p
            ciphertext.append(C2)

    return C1, ciphertext, space_positions

# Fungsi untuk dekripsi El Gamal
def decrypt(p, x, C1, ciphertext, space_positions):
    plaintext = []
    C1_x = pow(C1, x, p)
    C1_x_inv = pow(C1_x, -1, p)  # Invers modular

    char_index = 0
    for i in range(len(ciphertext) + len(space_positions)):
        if i in space_positions:
            plaintext.append(" ")
        else:
            C2 = ciphertext[char_index]
            M = (C2 * C1_x_inv) % p
            base = ord('A') if M < 26 else ord('a')
            plaintext.append(chr(M + base))
            char_index += 1

    return "".join(plaintext)

# Fungsi untuk membaca file DOCX
def read_docx(file_path):
    doc = Document(file_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text + '\n'
    return text.strip()

# Fungsi untuk menyimpan file DOCX
def save_docx(text, file_path):
    doc = Document()
    doc.add_paragraph(text)
    doc.save(file_path)

# Fungsi untuk membaca file TXT
def read_txt(file_path):
    with open(file_path, "r") as file:
        return file.read()

# Fungsi untuk menyimpan file TXT
def save_txt(text, file_path):
    with open(file_path, "w") as file:
        file.write(text)

# Streamlit UI
st.title("El Gamal Encryption for Document Protection")
mode = st.radio("Pilih Mode", ["Enkripsi", "Dekripsi"])
file_type = st.radio("Pilih Jenis File", ["DOCX", "TXT"])

if mode == "Enkripsi":
    st.subheader("Masukkan Parameter Kunci Publik")
    p = st.number_input("p (Bilangan Prima)", min_value=1, step=1)
    g = st.number_input("g (Generator)", min_value=1, step=1)
    x = st.number_input("x (Kunci Pribadi)", min_value=1, step=1)
    k = st.number_input("k (Bilangan Acak)", min_value=1, step=1)

    if all([p, g, x, k]):
        uploaded_file = st.file_uploader("Pilih file untuk dienkripsi", type=["docx", "txt"])

        if uploaded_file is not None:
            file_path = "temp_input_file." + file_type.lower()
            with open(file_path, "wb") as f:
                f.write(uploaded_file.getbuffer())

            if file_type == "DOCX":
                plaintext = read_docx(file_path)
            else:
                plaintext = read_txt(file_path)

            _, _, y, _ = generate_key(p, g, x)
            C1, ciphertext, space_positions = encrypt(p, g, y, k, plaintext)

            encrypted_text = f"C1: {C1}\nCiphertext: {' '.join(map(str, ciphertext))}\nSpaces: {' '.join(map(str, space_positions))}"
            st.text("File berhasil dienkripsi!")

            if file_type == "DOCX":
                save_docx(encrypted_text, "encrypted_file.docx")
                with open("encrypted_file.docx", "rb") as file:
                    st.download_button("Unduh File Terenkripsi", file, file_name="encrypted_file.docx")
            else:
                save_txt(encrypted_text, "encrypted_file.txt")
                with open("encrypted_file.txt", "rb") as file:
                    st.download_button("Unduh File Terenkripsi", file, file_name="encrypted_file.txt")

elif mode == "Dekripsi":
    st.subheader("Masukkan Parameter Kunci Pribadi")
    p = st.number_input("p (Bilangan Prima)", min_value=1, step=1)
    x = st.number_input("x (Kunci Pribadi)", min_value=1, step=1)

    if all([p, x]):
        uploaded_file = st.file_uploader("Pilih file hasil enkripsi", type=["docx", "txt"])

        if uploaded_file is not None:
            file_path = "temp_encrypted_file." + file_type.lower()
            with open(file_path, "wb") as f:
                f.write(uploaded_file.getbuffer())

            if file_type == "DOCX":
                encrypted_content = read_docx(file_path)
            else:
                encrypted_content = read_txt(file_path)

            lines = encrypted_content.split("\n")
            C1 = int(lines[0].split(": ")[1])
            ciphertext = list(map(int, lines[1].split(": ")[1].split()))
            space_positions = list(map(int, lines[2].split(": ")[1].split()))

            decrypted_text = decrypt(p, x, C1, ciphertext, space_positions)
            st.text("Hasil Dekripsi:")
            st.text(decrypted_text)

            if file_type == "DOCX":
                save_docx(decrypted_text, "decrypted_file.docx")
                with open("decrypted_file.docx", "rb") as file:
                    st.download_button("Unduh File Terdekripsi", file, file_name="decrypted_file.docx")
            else:
                save_txt(decrypted_text, "decrypted_file.txt")
                with open("decrypted_file.txt", "rb") as file:
                    st.download_button("Unduh File Terdekripsi", file, file_name="decrypted_file.txt")
